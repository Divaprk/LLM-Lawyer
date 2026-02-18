import os
from docx import Document

# Path to one of your files
BASE_DIR = os.path.dirname(os.path.dirname(__file__))
FILE_PATH = os.path.join(BASE_DIR, 'data', 'raw_data', 'employment_act.docx')

def inspect_file():
    if not os.path.exists(FILE_PATH):
        print(f"[ERROR] Could not find {FILE_PATH}")
        return

    doc = Document(FILE_PATH)
    
    print(f"--- INSPECTION REPORT FOR: {os.path.basename(FILE_PATH)} ---")
    print(f"Total Tables Found: {len(doc.tables)}")
    print(f"Total Paragraphs Found: {len(doc.paragraphs)}")
    
    # CHECK TABLES
    if len(doc.tables) > 0:
        print("\n[TABLE STRUCTURE DETECTED]")
        table = doc.tables[0]
        print(f"Table 1 has {len(table.rows)} rows and {len(table.columns)} columns.")
        print("--- First 3 Rows Content ---")
        for i, row in enumerate(table.rows[:3]):
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            print(f"Row {i}: {row_text}")
    else:
        print("\n[NO TABLES DETECTED] - The data is likely in Paragraphs.")

    # CHECK PARAGRAPHS
    print("\n--- First 10 Paragraphs Content ---")
    for i, p in enumerate(doc.paragraphs[:10]):
        if p.text.strip():
            print(f"P{i} [{p.style.name}]: {p.text[:100]}...")

if __name__ == "__main__":
    inspect_file()